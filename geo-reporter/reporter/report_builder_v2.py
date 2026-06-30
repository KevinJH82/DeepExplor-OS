"""
Report Builder V2 —— 融合论证式新版报告(Word),与旧版 report_builder.py **并存**。
旧版按数据类别罗列;新版按论证链组织(摘要→项目概况→地质背景→工作方法→数据综合分析→
靶区→钻孔方案→价值评估→风险→结论)。复用旧版 ReportBuilder 的样式/渲染 helper。

Phase A:摘要/项目概况/地质背景/工作方法/数据综合分析/结论 已用真实数据填充;
靶区/钻孔为现有数据直接列出;价值评估(Phase C)与拐点/预算(Phase B)暂为框架/占位。
"""
import math
from datetime import datetime
from typing import Dict, Optional

from docx import Document
from docx.shared import Cm

from .report_builder import ReportBuilder, _is_no_data
from .categories import SearchResult, get_all_categories
from .geocoder import LocationContext

# Phase B 参数
_TARGET_BOX_HALF_M = 150.0    # 靶区方框半边长(m):由靶点中心外扩成 ~300m×300m 圈定拐点
_DRILL_DEEP_M = 400.0         # 孔型阈值:目标深度 > 400m 用斜孔,否则直孔(参照 JORC 实践)
_DEFAULT_DRILL_RATE = 1500    # 钻探默认综合单价(元/m,占位):精确预算待 Phase C 经济参数录入


def _lonlat_to_utm(lon: float, lat: float):
    """WGS84 经纬度 → UTM(无第三方依赖)。返回 (zone, hemi, easting, northing)。"""
    a, f, k0 = 6378137.0, 1 / 298.257223563, 0.9996
    e2 = f * (2 - f)
    ep2 = e2 / (1 - e2)
    zone = int((lon + 180) / 6) + 1
    lon0 = math.radians((zone - 1) * 6 - 180 + 3)
    phi, lam = math.radians(lat), math.radians(lon)
    N = a / math.sqrt(1 - e2 * math.sin(phi) ** 2)
    T = math.tan(phi) ** 2
    C = ep2 * math.cos(phi) ** 2
    A = math.cos(phi) * (lam - lon0)
    M = a * ((1 - e2 / 4 - 3 * e2 ** 2 / 64 - 5 * e2 ** 3 / 256) * phi
             - (3 * e2 / 8 + 3 * e2 ** 2 / 32 + 45 * e2 ** 3 / 1024) * math.sin(2 * phi)
             + (15 * e2 ** 2 / 256 + 45 * e2 ** 3 / 1024) * math.sin(4 * phi)
             - (35 * e2 ** 3 / 3072) * math.sin(6 * phi))
    easting = k0 * N * (A + (1 - T + C) * A ** 3 / 6
                        + (5 - 18 * T + T ** 2 + 72 * C - 58 * ep2) * A ** 5 / 120) + 500000
    northing = k0 * (M + N * math.tan(phi) * (A ** 2 / 2 + (5 - T + 9 * C + 4 * C ** 2) * A ** 4 / 24
                     + (61 - 58 * T + T ** 2 + 600 * C - 330 * ep2) * A ** 6 / 720))
    if lat < 0:
        northing += 10000000.0
    return zone, ("S" if lat < 0 else "N"), easting, northing

# 论证式骨架(供 PPT v2 共用骨架顺序)
_V2_CHAPTERS = [
    ("摘要", "核心发现、建议钻孔方案、资源潜力与价值、风险与建议(执行摘要)。"),
    ("项目概况", "项目来源、目的任务、勘查区交通位置及范围、取得的主要成果。"),
    ("区域与矿区地质背景", "区域地质构造与成矿规律;矿区地层、侵入岩、构造、矿体特征与围岩蚀变。"),
    ("工作方法与技术", "多模态异构数据 AI 智能找矿技术;靶向超弱核磁共振 AI 探测(预留);数据源与证据链。"),
    ("数据综合分析", "物探、化探、遥感、InSAR 等证据及其相互印证;已知矿点/钻孔验证。"),
    ("隐伏矿位置评估", "靶区圈定方法(地-空-钻三位一体);各靶区位置依据、拐点坐标与预测参数。"),
    ("建议钻孔方案", "钻孔设计原则;孔位表(拐点坐标、孔深、孔型、优先级)与预算。"),
    ("资源潜力与价值评估", "资源增量情景(保守/基准/乐观)、koz、ROI 与发现成本(需录入经济参数)。"),
    ("风险评估与建议", "技术/工程/经济风险及缓解措施;行动建议。"),
    ("结论与建议", "综合结论与下一步工作建议。"),
    ("参考资料", "文献与数据来源清单。"),
]


class ReportBuilderV2(ReportBuilder):
    """融合论证式新版报告生成器(复用旧版样式/渲染 helper)。"""

    def build_report_v2(self, location: LocationContext, search_results: Dict[str, SearchResult],
                        output_name: Optional[str] = None, mineral_type: str = "",
                        target_figure=None, confidence: dict = None, tenant_id: str = None,
                        econ_params: dict = None) -> str:
        from .data_sources import (set_tenant, fetch_datacolle_metallogenic,
                                    fetch_datacolle_literature, fetch_datacolle_papers)
        set_tenant(tenant_id)
        bbox = (location.min_lon, location.min_lat, location.max_lon, location.max_lat)
        dc_met = fetch_datacolle_metallogenic(*bbox)
        dc_lit = fetch_datacolle_literature(*bbox)
        dc_pap = fetch_datacolle_papers(*bbox)

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2.54)
            sec.left_margin = sec.right_margin = Cm(3.17)

        area = (getattr(location, "area_name", "") or "").strip()
        if not area or area.lower() == "aoi":
            area = (getattr(location, "location_str", "") or "").strip() or "研究区"
        # v2 末位生成,统一用解析后的真实区域名(避免摘要等复用 helper 显示占位 "AOI")
        try:
            location.area_name = area
        except Exception:
            pass

        # ===== 封面 =====
        self._add_paragraph(doc, "")
        self._add_heading(doc, f"{area} 地质勘探综合评估报告(融合版)", level=1)
        cover = [f"位置：{location.location_str}（{location.country}）",
                 f"坐标范围：{location.coords_str}"]
        if mineral_type:
            cover.append(f"目标矿种：{mineral_type}")
        cover += [f"报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}"]
        self._add_paragraph(doc, "\n".join(cover), font_size=12)
        self._add_paragraph(doc, "【新版·融合论证式】本版按「项目概况→地质背景→工作方法→数据综合分析→"
                                 "靶区→钻孔方案→价值评估→风险→结论」论证链组织,与旧版报告并存提供。", font_size=10)
        doc.add_page_break()

        ch = [0]  # 章节计数器(可变,供闭包递增)

        def chap(title):
            ch[0] += 1
            self._add_heading(doc, f"第{self._num_to_chinese(ch[0])}章  {title}", level=1)

        # ===== 摘要(执行摘要)=====
        self._add_heading(doc, "摘要", level=1)
        try:
            self._add_front_matter(doc, location, search_results, mineral_type,
                                   target_figure=target_figure, confidence=confidence)
        except Exception as e:
            self._add_paragraph(doc, f"(执行摘要生成异常,详见各章)：{e}", font_size=10)
        doc.add_page_break()

        # ===== 第1章 项目概况 =====
        chap("项目概况")
        self._add_heading(doc, "1.1 勘查区位置与范围", level=2)
        self._add_paragraph(doc, f"研究区位于 {location.location_str}（{location.country}），"
                                 f"坐标范围为 {location.coords_str}。"
                                 + (f"目标矿种为{mineral_type}。" if mineral_type else ""))
        _admin = [(k, v) for k, v in (("国家", location.country), ("省份", location.province),
                                      ("城市", location.city), ("区县", location.district)) if v and str(v).strip()]
        if _admin:
            self._add_paragraph(doc, "行政区划：" + "；".join(f"{k}{v}" for k, v in _admin) + "。")
        self._add_heading(doc, "1.2 取得的主要成果", level=2)
        self._add_paragraph(doc, self._v2_achievements(search_results, target_figure, confidence, mineral_type))
        doc.add_page_break()

        # ===== 第2章 区域与矿区地质背景 =====
        chap("区域与矿区地质背景")
        geo = search_results.get("geology")
        if geo and not getattr(geo, "error", None):
            if geo.summary and not _is_no_data(geo.summary):
                self._add_paragraph(doc, geo.summary, font_size=11)
            dps = [dp for dp in (geo.data_points or []) if not _is_no_data(dp.value)]
            if dps:
                self._add_table(doc, ["项目", "数值/描述"], [[dp.item, dp.value] for dp in dps], col_widths=[5, 9])
            if geo.key_findings:
                self._add_key_findings(doc, geo.key_findings)
        else:
            self._add_paragraph(doc, "本区地质资料待补充。", font_size=11)
        self._add_metallogenic_evidence(doc, dc_met, dc_lit, dc_pap)
        doc.add_page_break()

        # ===== 第3章 工作方法与技术 =====
        chap("工作方法与技术")
        self._add_heading(doc, "3.1 多模态异构数据 AI 智能找矿技术", level=2)
        self._add_paragraph(doc, "本次评估采用多模态异构数据 AI 智能找矿技术,融合遥感蚀变、构造解译、"
                                 "地球物理(磁/重)、地球化学与 InSAR 形变等多源证据,经三维成矿建模(geo-model3d)"
                                 "形成深部有利度体与靶点,构成「地-空-钻」三位一体的证据链。", font_size=11)
        self._add_heading(doc, "3.2 靶向超弱核磁共振 AI 探测技术(预留)", level=2)
        self._add_paragraph(doc, "靶向超弱核磁共振 AI 探测为本平台预留能力,可对重点靶区开展直接探测验证;"
                                 "本次报告暂未接入该技术实测数据。", font_size=11)
        self._add_heading(doc, "3.3 数据源与证据链", level=2)
        rows = []
        for cat in get_all_categories():
            if cat.id == "slow_variables":
                continue
            r = search_results.get(cat.id)
            if r and not getattr(r, "error", None):
                rows.append([cat.chapter_title, getattr(r, "evidence_level", "") or "—",
                             "已接入" if (r.summary or r.data_points or r.key_findings) else "占位"])
        if rows:
            self._add_table(doc, ["数据类别", "证据来源层级", "状态"], rows, col_widths=[5, 6, 3])
        doc.add_page_break()

        # ===== 第4章 数据综合分析 =====
        chap("数据综合分析")
        self._add_paragraph(doc, "下列各类证据按「是否与构造-蚀变-异常在空间上同向叠合」进行综合判读;"
                                 "单项异常仅在与其他证据交汇时方上升为靶区线索。", font_size=11)
        sub = 0
        for cat in get_all_categories():
            if cat.id in ("geology", "slow_variables"):
                continue
            r = search_results.get(cat.id)
            if not r or getattr(r, "error", None):
                continue
            if not (r.summary or r.data_points or r.key_findings):
                continue
            sub += 1
            self._add_heading(doc, f"4.{sub} {cat.chapter_title}", level=2)
            if r.summary and not _is_no_data(r.summary):
                self._add_paragraph(doc, r.summary, font_size=11)
            dps = [dp for dp in (r.data_points or []) if not _is_no_data(dp.value)]
            if dps:
                self._add_table(doc, ["项目", "数值/描述"], [[dp.item, dp.value] for dp in dps], col_widths=[5, 9])
            if r.key_findings:
                self._add_key_findings(doc, r.key_findings)
            for fig in (getattr(r, "figures", None) or []):
                self._add_figure(doc, fig)
        if sub == 0:
            self._add_paragraph(doc, "本次暂无可纳入综合分析的物化探/遥感/形变证据。", font_size=11)
        doc.add_page_break()

        # ===== 第5章 隐伏矿位置评估(靶区)=====
        chap("隐伏矿位置评估")
        self._v2_targets(doc, target_figure)
        doc.add_page_break()

        # ===== 第6章 建议钻孔方案 =====
        chap("建议钻孔方案")
        self._v2_drill(doc, location, target_figure)
        doc.add_page_break()

        # ===== 第7章 资源潜力与价值评估(参数化:上传经济参数表后出定量)=====
        chap("资源潜力与价值评估")
        self._add_paragraph(doc, self._v2_resource_qualitative(target_figure), font_size=11)
        self._v2_value_assessment(doc, econ_params, target_figure, location)
        doc.add_page_break()

        # ===== 第8章 风险评估与建议 =====
        chap("风险评估与建议")
        self._v2_risk(doc, search_results, confidence)
        doc.add_page_break()

        # ===== 第9章 结论与建议 =====
        chap("结论与建议")
        self._add_confidence_section(doc, confidence, mineral_type)

        # ===== 参考资料 =====
        self._add_heading(doc, "参考资料", level=1)
        papers = dc_pap if isinstance(dc_pap, list) else (dc_pap or {}).get("papers", []) if dc_pap else []
        if papers:
            for i, p in enumerate(papers[:30], 1):
                if isinstance(p, dict):
                    line = f"[{i}] " + "，".join(str(x) for x in (p.get("title"), p.get("author") or (p.get("authors") or [None])[0], p.get("year"), p.get("doi") or p.get("url")) if x)
                else:
                    line = f"[{i}] {p}"
                self._add_paragraph(doc, line, font_size=9)
        else:
            self._add_paragraph(doc, "本次未检索到可引用的公开文献。", font_size=10)

        if output_name is None:
            output_name = f"{area}_{datetime.now().strftime('%Y%m%d')}"
        out = self.output_dir / f"{output_name}_新版.docx"
        doc.save(str(out))
        return str(out)

    # ---- v2 专用小工具 ----
    def _v2_achievements(self, search_results, target_figure, confidence, mineral_type) -> str:
        n_done = sum(1 for r in (search_results or {}).values() if r and not getattr(r, "error", None))
        targets = getattr(target_figure, "targets", None) or []
        grade = (confidence or {}).get("grade") or (confidence or {}).get("overall_grade")
        parts = [f"本次综合 {n_done} 类地学资料,经多源证据融合与三维成矿建模"]
        if targets:
            depths = [t.get("target_depth_m") for t in targets if t.get("target_depth_m") not in (None, "")]
            dtxt = f"、目标深度约 {min(depths)}–{max(depths)} m" if depths else ""
            parts.append(f"圈定 {len(targets)} 个三维预测靶区{dtxt}")
        if grade:
            parts.append(f"综合成矿置信定级为 {grade} 级")
        tail = "。" + (f"目标矿种为{mineral_type}。" if mineral_type else "")
        return "，".join(parts) + tail

    def _v2_targets(self, doc, target_figure):
        if target_figure is None:
            self._add_paragraph(doc, "本研究区暂无三维成矿预测靶点,靶区结论待补充。", font_size=11)
            return
        self._add_paragraph(doc, "采用 geo-model3d 三维成矿建模的 targets_3d 作为靶区来源,"
                                 "按「地-空-钻」三位一体圈定隐伏矿有利位置,并给出 A–D 置信评级。", font_size=11)
        try:
            self._add_figure(doc, target_figure, width_cm=15.0)
        except Exception:
            pass
        targets = getattr(target_figure, "targets", None) or []
        if not targets:
            return
        rows = [[f"#{t.get('rank','')}", t.get("grade", ""),
                 f"{t.get('longitude',0):.4f}, {t.get('latitude',0):.4f}",
                 str(t.get("target_depth_m", "")), t.get("reason", "")] for t in targets]
        self._add_table(doc, ["靶区", "置信等级", "中心坐标(°E, °N)", "目标深度(m)", "评分理由"],
                        rows, col_widths=[1.4, 1.7, 3.4, 2.0, 7.0])

        # 各靶区拐点坐标(中心外扩 ~300m 方框,出 UTM 与经纬度;取置信靠前的前 5 个)
        self._add_paragraph(doc, "")
        self._add_paragraph(doc, f"各靶区拐点坐标(以靶点中心外扩 {int(_TARGET_BOX_HALF_M*2)} m 方框圈定,"
                                 "UTM/WGS84 并列,供野外查证):", font_size=11)
        letters = "ABCDE"
        for i, t in enumerate(targets[:5]):
            lon0, lat0 = float(t.get("longitude", 0) or 0), float(t.get("latitude", 0) or 0)
            if not lon0 and not lat0:
                continue
            dlat = _TARGET_BOX_HALF_M / 111320.0
            dlon = _TARGET_BOX_HALF_M / (111320.0 * max(math.cos(math.radians(lat0)), 1e-6))
            corners = [("拐点1(西北)", lon0 - dlon, lat0 + dlat), ("拐点2(东北)", lon0 + dlon, lat0 + dlat),
                       ("拐点3(东南)", lon0 + dlon, lat0 - dlat), ("拐点4(西南)", lon0 - dlon, lat0 - dlat)]
            zone, hemi, _, _ = _lonlat_to_utm(lon0, lat0)
            crows = []
            for name, clon, clat in corners:
                _, _, e, n = _lonlat_to_utm(clon, clat)
                crows.append([name, f"{e:,.0f}", f"{n:,.0f}", f"{clon:.5f}", f"{clat:.5f}"])
            self._add_paragraph(doc, f"靶区 {letters[i]}(对应 #{t.get('rank','')},置信 {t.get('grade','')},"
                                     f"目标深度 {t.get('target_depth_m','')} m;UTM {zone}{hemi} 带):", font_size=10)
            self._add_table(doc, ["拐点", "UTM东(m)", "UTM北(m)", "经度(°E)", "纬度(°N)"], crows,
                            col_widths=[2.6, 2.8, 2.8, 2.9, 2.9])

    def _v2_drill(self, doc, location, target_figure):
        try:
            from .synthesis import get_drill_evidence
            drill = get_drill_evidence(location) or {}
        except Exception:
            drill = {}
        holes = drill.get("holes") or []
        feedback = drill.get("feedback") or []
        if not holes and not feedback:
            self._add_paragraph(doc, "本次暂无 geo-drill 布孔/钻探产物;建议在靶区优选后开展钻探设计。", font_size=11)
            return
        self._add_paragraph(doc, "6.1 钻孔设计原则", font_size=11, bold=True)
        self._add_paragraph(doc, f"按目标深度选孔型:目标深度 ≤ {int(_DRILL_DEEP_M)} m 用直孔(近垂直矿体,"
                                 f"施工效率高);> {int(_DRILL_DEEP_M)} m 用斜孔(验证深部及陡倾矿体)。"
                                 "孔位取自 geo-drill 决策支持,坐标以 UTM/WGS84 并列。", font_size=11)
        total_m = 0.0
        if holes:
            rows = []
            for h in holes[:20]:
                lon, lat = float(h.get("lon", 0) or 0), float(h.get("lat", 0) or 0)
                dep = h.get("target_depth_m", "")
                try:
                    depf = float(dep); total_m += depf
                except (TypeError, ValueError):
                    depf = 0.0
                kind = "斜孔" if depf > _DRILL_DEEP_M else "直孔"
                z, hemi, e, n = _lonlat_to_utm(lon, lat) if (lon or lat) else (0, "", 0, 0)
                rows.append([h.get("hole_id", "") or f"#{h.get('rank','')}", kind,
                             f"{e:,.0f}", f"{n:,.0f}", f"{lon:.4f}, {lat:.4f}",
                             str(dep), str(h.get("priority", h.get("score", "")))])
            self._add_paragraph(doc, "6.2 孔位与孔型", font_size=11, bold=True)
            self._add_table(doc, ["计划孔", "孔型", "UTM东(m)", "UTM北(m)", "经纬度(°E,°N)", "目标深度(m)", "优先级/评分"],
                            rows, col_widths=[2.0, 1.4, 2.6, 2.6, 3.0, 2.2, 2.0])
        # 预算明细(默认单价占位,精确预算待 Phase C 经济参数录入)
        self._add_paragraph(doc, "6.3 预算估算(默认单价占位)", font_size=11, bold=True)
        n_hole = len(holes[:20])
        drill_cost = total_m * _DEFAULT_DRILL_RATE
        brows = [
            ["钻探进尺", f"{total_m:,.0f} m", f"{_DEFAULT_DRILL_RATE:,} 元/m", f"{drill_cost:,.0f} 元"],
            ["动复员/场地", f"{n_hole} 孔", "—", "(待录入)"],
            ["岩矿测试化验", f"{n_hole} 孔", "—", "(待录入)"],
            ["管理与编录", "—", "—", "(待录入)"],
        ]
        self._add_table(doc, ["预算项", "数量", "单价", "小计"], brows, col_widths=[3.5, 3.0, 3.5, 4.0])
        self._add_paragraph(doc, f"共 {n_hole} 孔、总进尺约 {total_m:,.0f} m;以上仅按默认综合单价"
                                 f"{_DEFAULT_DRILL_RATE:,} 元/m 估算钻探进尺费,动复员/化验/管理及精确单价"
                                 "待 Phase C 经济参数录入后核定。", font_size=9)
        if feedback:
            ore = sum(1 for f in feedback if f.get("outcome") == "ore")
            self._add_paragraph(doc, f"已有钻孔反馈 {len(feedback)} 个:见矿 {ore},无矿 {len(feedback)-ore}。", font_size=10)

    def _v2_drill_total_m(self, location) -> float:
        """geo-drill 计划孔总进尺(m),用于价值评估预算估算;取不到返回 0。"""
        if location is None:
            return 0.0
        try:
            from .synthesis import get_drill_evidence
            holes = (get_drill_evidence(location) or {}).get("holes") or []
        except Exception:
            return 0.0
        tot = 0.0
        for h in holes:
            try:
                tot += float(h.get("target_depth_m") or 0)
            except (TypeError, ValueError):
                pass
        return tot

    def _v2_resource_qualitative(self, target_figure) -> str:
        targets = getattr(target_figure, "targets", None) or []
        if not targets:
            return "本区资源潜力评估有待靶区与钻探数据支撑。"
        depths = [t.get("target_depth_m") for t in targets if t.get("target_depth_m") not in (None, "")]
        dtxt = f",目标深度约 {min(depths)}–{max(depths)} m" if depths else ""
        return (f"三维建模圈定 {len(targets)} 个预测靶区{dtxt},指示本区具备一定隐伏矿资源潜力;"
                f"靶区沿主控构造延伸方向展布,深部仍有扩边空间。")

    def _v2_value_assessment(self, doc, econ_params, target_figure, location=None):
        """⑦章定量:有上传经济参数表 → 三情景资源增量/价值/ROI/发现成本;否则提示上传参数表。
        预算进尺优先用 geo-drill 钻孔总进尺(更贴近实际),无则退化为靶点深度之和。"""
        from .value_assessment import compute_value_assessment
        targets = getattr(target_figure, "targets", None) or []
        total_m = self._v2_drill_total_m(location)
        if not total_m:
            for t in targets:
                try:
                    total_m += float(t.get("target_depth_m") or 0)
                except (TypeError, ValueError):
                    pass
        va = compute_value_assessment(econ_params or {}, targets, total_m)
        if not va:
            self._add_paragraph(doc, "定量价值评估(资源增量情景、koz、ROI、发现成本)需上传经济参数表"
                                     "(现有资源量/品位、金属价格、钻探单价等);上传后本章将自动给出"
                                     "保守/基准/乐观三情景估算。", font_size=10)
            return
        cur = va["currency"]
        self._add_paragraph(doc, f"基于上传经济参数(现有资源约 {va['existing_koz']:,.0f} koz,金属价格 "
                                 f"{va['metal_price']:,.0f} {cur}/oz"
                                 + (f",勘探预算约 {va['budget']:,.0f} {cur}" if va.get("budget") else "")
                                 + "),给出三情景资源增量与价值评估:", font_size=11)
        rows = []
        for s in va["scenarios"]:
            rows.append([
                f"{s['name']}({s['increment_pct']:.0f}%)",
                f"{s['new_koz']:,.0f}",
                f"{s['total_koz']:,.0f}",
                f"{s['new_value']/1e6:,.1f}",
                (f"{s['roi_pct']:,.0f}%" if s['roi_pct'] is not None else "—"),
                (f"{s['discovery_cost_per_oz']:,.0f}" if s['discovery_cost_per_oz'] is not None else "—"),
            ])
        self._add_table(doc, ["情景", "新增(koz)", "总资源(koz)", f"新增价值(百万{cur})", "ROI", f"发现成本({cur}/oz)"],
                        rows, col_widths=[2.6, 2.2, 2.4, 2.8, 2.2, 2.6])
        self._add_paragraph(doc, "注:资源增量为靶区相对现有资源的比例情景估算;ROI=新增价值/勘探预算,"
                                 "发现成本=勘探预算/新增金属盎司。数值随上传参数而定,供投资决策参考。", font_size=9)

    def _v2_risk(self, doc, search_results, confidence):
        failed = [c.chapter_title for c in get_all_categories()
                  if (search_results or {}).get(c.id) and getattr(search_results[c.id], "error", None)]
        self._add_paragraph(doc, "技术风险:三维预测为多源证据综合推断,深部矿化连续性需钻探直接验证;"
                                 "建议对高置信靶区优先加密验证。", font_size=11)
        if failed:
            self._add_paragraph(doc, "数据风险:以下证据本次缺失或获取失败,可能影响判别充分性——"
                                     + "、".join(failed) + "。", font_size=11)
        self._add_paragraph(doc, "工程与经济风险:深孔施工(涌水等)与金属价格波动需在实施阶段评估;"
                                 "相关定量分析待经济参数录入后给出。", font_size=11)
