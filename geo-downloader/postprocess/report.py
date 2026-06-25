"""
Postprocess: 卫星数据下载报告生成模块
在每次打包完成后自动在 delivery/{area}/ 目录生成 Word 报告。

报告文件名：{area_label}_卫星数据下载报告_{YYYYMMDD}.docx

报告结构：
  一、任务概览（表格：任务时间、传感器数量、交付包大小等）
  二、各传感器下载明细（表格：传感器、类型、文件数、数据量）
  三、未获取数据传感器（表格：传感器、类型、原因）
  四、交付成果（表格：夏季包、冬季包内容、输出路径）
  五、备注（分辨率提升说明等）

依赖：python-docx（可选，未安装则跳过报告生成）
"""

from __future__ import annotations

import math
import os
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Tuple


# ── 传感器元数据（用于报告中的类型、分辨率描述）─────────────────────
_SENSOR_META = {
    "sentinel2":    ("多光谱",   "Sentinel-2 L2A",      "10-60m"),
    "sentinel1":    ("SAR 雷达", "Sentinel-1 GRD",      "10m"),
    "landsat":      ("多光谱",   "Landsat 8/9 L2",      "30m→15m"),
    "landsat7":     ("多光谱",   "Landsat 7 ETM+",      "15-30m"),
    "landsat_tirs": ("热红外",   "Landsat TIRS",        "30m"),
    "emit":         ("高光谱",   "EMIT L2A",            "60m→30m"),
    "dem":          ("高程",     "Copernicus DEM GLO-30", "30m→15m"),
    "srtm":         ("高程",     "SRTM DEM",              "30m→15m"),
    "aster":        ("热红外",   "ASTER L2",            "15-90m"),
    "aster_l1t":    ("热红外",   "ASTER L1T",           "15-90m"),
    "modis":        ("多光谱",   "MODIS",               "250-500m"),
    "alos":         ("SAR 雷达", "ALOS PALSAR",         "—"),
    "alos2":        ("SAR 雷达", "ALOS-2 PALSAR-2",     "1-100m"),
    "gedi":         ("激光雷达", "GEDI L2A",            "25m"),
    "opera":        ("SAR 雷达", "OPERA RTC-S1",        "30m"),
    "ecostress":    ("热红外",   "ECOSTRESS",           "70m"),
    "enmap":        ("高光谱",   "EnMAP L2A",           "30m"),
    "hyperion":     ("高光谱",   "Hyperion L1",         "30m"),
    "aviris":       ("高光谱",   "AVIRIS-NG",           "~5m"),
    "planet":       ("多光谱",   "PlanetScope",         "3-5m"),
    "prisma":       ("高光谱",   "PRISMA L2D",          "30m"),
    "desis":        ("高光谱",   "DESIS L2A",           "30m"),
    "zy1":          ("高光谱",   "ZY-1 02D AHSI",       "30m"),
    "spot67":       ("多光谱",   "SPOT 6/7",            "1.5-6m"),
    "pleiades":     ("多光谱",   "Pleiades 1A/1B",      "0.5-2m"),
    "wv2":          ("多光谱",   "WorldView-2",         "0.46-1.85m"),
    "wv3":          ("多光谱",   "WorldView-3",         "0.31-1.24m"),
    "nisar":        ("SAR 雷达", "NISAR",               "3-25m"),
}

# 兜底文案：只在动态判断无法确定真实状态时使用
_SENSOR_NO_DATA_REASON_FALLBACK = {
    "aviris":  "无覆盖该区域的存档数据",
    "desis":   "无覆盖该区域的存档数据",
    "alos2":   "无可访问数据源",
    "nisar":   "无可访问数据源",
    "enmap":   "需手动申请（DLR EOWEB）",
    "zy1":     "需手动申请（CRESDA）",
    "prisma":  "需手动申请（ASI）",
}


def _resolve_no_data_reason(sensor: str, raw_area_dir: Optional[Path]) -> str:
    """
    根据 raw 状态文件动态判断"未获取数据"的真实原因。
    Why: EnMAP/PRISMA 的硬编码"需手动申请"具有误导性 —— 代码其实有完整的自动化流程，
         返回 0 的真实原因可能是 (a) 该区域确实无存档数据，或 (b) 订单已自动提交但 ASI 还在排队。
    raw_area_dir : downloads/{area}/ 那一层（不是 delivery_dir）。
    """
    if raw_area_dir is not None:
        try:
            raw_area_dir = Path(raw_area_dir)
        except Exception:
            raw_area_dir = None
    if raw_area_dir and raw_area_dir.exists():
        # PRISMA: 有 .prisma_pending_order.json 说明订单已自动提交
        if sensor == "prisma":
            pending = raw_area_dir / "prisma" / ".prisma_pending_order.json"
            if pending.exists():
                return "订单已自动提交至 ASI，处理中（下次任务自动续传）"
        # EnMAP: enmap_debug.log 出现"找到 0 景"说明该 ROI/时段确实无归档
        # 日志位置：raw_area_dir 是 downloads/{area}/，enmap log 在同级 downloads/enmap/
        if sensor == "enmap":
            for candidate in (
                raw_area_dir.parent / "enmap" / "enmap_debug.log",
                raw_area_dir / "enmap" / "enmap_debug.log",
            ):
                if candidate.exists():
                    try:
                        txt = candidate.read_text(errors="ignore")
                    except Exception:
                        continue
                    if "找到 0 景" in txt:
                        return "该区域 DLR EOWEB 无存档数据"
                    break
    return _SENSOR_NO_DATA_REASON_FALLBACK.get(sensor, "搜索范围内无符合条件的景次")

_SEASON_SUMMER_LABEL = "data-矿权-夏季（6-8月）"
_SEASON_WINTER_LABEL = "data-矿权-冬季（11-3月）"


def _dir_size_mb(path: Path) -> float:
    """递归计算目录总大小（MB）"""
    total = 0
    if path.exists():
        for f in path.rglob("*"):
            if f.is_file():
                try:
                    total += f.stat().st_size
                except OSError:
                    pass
    return total / (1024 * 1024)


def _human_size(mb: float) -> str:
    if mb >= 1024:
        return f"{mb / 1024:.1f} GB"
    if mb >= 1:
        return f"{mb:.0f} MB"
    return f"{mb * 1024:.0f} KB"


def _count_files_in_season(season_dir: Path) -> Dict[str, int]:
    """统计 season_dir 下各传感器子目录的文件数"""
    counts: Dict[str, int] = {}
    if not season_dir.exists():
        return counts
    for entry in season_dir.iterdir():
        if entry.is_dir():
            n = sum(1 for f in entry.rglob("*") if f.is_file())
            counts[entry.name] = n
        elif entry.is_file() and entry.suffix.lower() in (".tif", ".tiff", ".nc", ".hdf"):
            counts[entry.name] = 1
    return counts


def _season_summary(delivery_dir: Path) -> Tuple[str, str]:
    """生成夏季/冬季交付内容的简要描述字符串"""
    def describe(season_dir: Path) -> str:
        if not season_dir.exists():
            return "（无数据）"
        parts = []
        for entry in sorted(season_dir.iterdir()):
            if entry.is_dir():
                parts.append(entry.name)
            elif entry.is_file() and entry.suffix.lower() in (".tif", ".tiff"):
                parts.append(entry.name)
        return "、".join(parts) if parts else "（无数据）"

    summer = describe(delivery_dir / _SEASON_SUMMER_LABEL)
    winter = describe(delivery_dir / _SEASON_WINTER_LABEL)
    return summer, winter


def _set_cell_bold(cell, bold: bool = True):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
        if not para.runs and para.text:
            run = para.add_run(para.text)
            para.clear()
            para.add_run(para.text).bold = bold


def _add_table_header(table, headers: List[str]):
    """设置表格第一行为加粗表头"""
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def generate_report(
    delivery_dir: Path,
    area_label: str,
    sensors_attempted: List[str],
    summary: Dict[str, int],
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    start_time: Optional[datetime] = None,
    end_time: Optional[datetime] = None,
    raw_area_dir: Optional[Path] = None,
) -> Optional[Path]:
    """
    生成卫星数据下载报告 Word 文档，存入 delivery_dir。

    Parameters
    ----------
    delivery_dir      : 交付目录路径（delivery/{area}/）
    area_label        : 区域名称
    sensors_attempted : 本次尝试下载的传感器列表
    summary           : {sensor_id: file_count} 下载结果统计
    start_date        : 搜索起始日期字符串（可选）
    end_date          : 搜索结束日期字符串（可选）
    start_time        : 任务开始时间（可选，用于计算耗时）
    end_time          : 任务结束时间（可选）

    Returns
    -------
    Path - 生成的 docx 路径；若 python-docx 未安装则返回 None
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  [报告] 未安装 python-docx，跳过报告生成（pip install python-docx）")
        return None

    now = datetime.now()
    report_name = f"{area_label}_卫星数据下载报告_{now.strftime('%Y%m%d')}.docx"
    report_path = delivery_dir / report_name

    # 若已存在同名报告则覆盖（重新打包时更新）
    doc = Document()

    # ── 全局字体设置 ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ── 标题 ─────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{area_label}  卫星数据下载报告")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"任务日期：{now.strftime('%Y-%m-%d')}")

    doc.add_paragraph()  # 空行

    # ── 一、任务概览 ──────────────────────────────────────────────
    doc.add_heading("一、任务概览", level=2)

    # 计算统计数据
    sensors_ok = [s for s in sensors_attempted
                  if s not in ("derive", "delivery") and summary.get(s, 0) > 0]
    sensors_fail = [s for s in sensors_attempted
                    if s not in ("derive", "delivery") and summary.get(s, 0) == 0]
    total_files = sum(v for k, v in summary.items() if k not in ("derive", "delivery"))
    delivery_mb = _dir_size_mb(delivery_dir)

    # 耗时
    if start_time and end_time:
        elapsed = end_time - start_time
        total_sec = int(elapsed.total_seconds())
        h, rem = divmod(total_sec, 3600)
        m = rem // 60
        duration_str = f"{h} 小时 {m} 分钟" if h > 0 else f"{m} 分钟"
        start_str = start_time.strftime("%H:%M")
        end_str = end_time.strftime("%H:%M")
    else:
        duration_str = "—"
        start_str = "—"
        end_str = "—"

    date_range = f"{start_date} ~ {end_date}" if start_date and end_date else "—"

    overview_rows = [
        ("任务开始",    start_str),
        ("任务结束",    end_str),
        ("总耗时",      duration_str),
        ("搜索时间范围", date_range),
        ("下载文件数",  f"{total_files} 个（原始文件）"),
        ("成功传感器",  f"{len(sensors_ok)} 种"),
        ("交付包大小",  _human_size(delivery_mb) + "（裁切 / 整理后）"),
    ]

    tbl1 = doc.add_table(rows=len(overview_rows), cols=2)
    tbl1.style = "Table Grid"
    for i, (k, v) in enumerate(overview_rows):
        row = tbl1.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    # ── 二、各传感器下载明细 ───────────────────────────────────────
    doc.add_heading("二、各传感器下载明细", level=2)

    if sensors_ok:
        tbl2 = doc.add_table(rows=1 + len(sensors_ok), cols=4)
        tbl2.style = "Table Grid"
        _add_table_header(tbl2, ["传感器", "类型", "文件数", "分辨率"])
        for i, sensor in enumerate(sensors_ok, start=1):
            meta = _SENSOR_META.get(sensor, ("—", sensor, "—"))
            row = tbl2.rows[i]
            row.cells[0].text = meta[1]   # 全称
            row.cells[1].text = meta[0]   # 类型
            row.cells[2].text = str(summary.get(sensor, 0))
            row.cells[3].text = meta[2]   # 分辨率
    else:
        doc.add_paragraph("（本次无传感器成功下载数据）")

    doc.add_paragraph()

    # ── 三、未获取数据传感器 ───────────────────────────────────────
    doc.add_heading("三、未获取数据传感器", level=2)

    if sensors_fail:
        tbl3 = doc.add_table(rows=1 + len(sensors_fail), cols=3)
        tbl3.style = "Table Grid"
        _add_table_header(tbl3, ["传感器", "类型", "原因"])
        for i, sensor in enumerate(sensors_fail, start=1):
            meta = _SENSOR_META.get(sensor, ("—", sensor, "—"))
            reason = _resolve_no_data_reason(sensor, raw_area_dir)
            row = tbl3.rows[i]
            row.cells[0].text = meta[1]
            row.cells[1].text = meta[0]
            row.cells[2].text = reason
    else:
        doc.add_paragraph("（所有传感器均获取到数据）")

    doc.add_paragraph()

    # ── 四、交付成果 ──────────────────────────────────────────────
    doc.add_heading("四、交付成果", level=2)

    summer_desc, winter_desc = _season_summary(delivery_dir)
    delivery_rows = [
        ("夏季包（6–8月）",   summer_desc),
        ("冬季包（11–3月）",  winter_desc),
        ("交付包总大小",      _human_size(delivery_mb) + "（已完成 ROI 裁切、波段整理）"),
        ("输出路径",          str(delivery_dir)),
    ]

    tbl4 = doc.add_table(rows=len(delivery_rows), cols=2)
    tbl4.style = "Table Grid"
    for i, (k, v) in enumerate(delivery_rows):
        row = tbl4.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    # ── 五、分辨率提升说明 ─────────────────────────────────────────
    doc.add_heading("五、分辨率说明", level=2)

    notes = [
        "本次交付已对低分辨率波段进行重采样以提升视觉质量：",
        "• Sentinel-2 B01/B05-B09（原 20-60m）→ 统一重采样至 10m，与其他波段对齐",
        "• Landsat 8/9 各波段（原 30m）→ 重采样至 15m（双线性插值）",
        "• Landsat 7 多光谱波段（原 30m）→ 重采样至 15m；全色波段 B8 保持原始 15m",
        "• EMIT 高光谱（原 60m）→ 重采样至 30m（双线性插值，提升 GIS 叠合精度）",
        "• DEM（Copernicus GLO-30，原 30m）→ 重采样至 15m，与 Landsat 交付分辨率对齐",
        "注意：重采样使用双线性插值，不增加实际信息量，仅提升像素密度与显示效果。",
    ]
    for note in notes:
        p = doc.add_paragraph(note)
        p.paragraph_format.space_after = Pt(2)

    # ── 六、数据质量分析 ──────────────────────────────────────────
    doc.add_heading("六、数据质量分析", level=2)
    try:
        from postprocess.quality_report import analyze_delivery, summarize
        qr_results = analyze_delivery(delivery_dir)
        qsum = summarize(qr_results)

        # 概览行
        ok_files = qsum["total_files"] - qsum["files_with_issues"]
        overview_q = [
            ("分析文件总数",   f"{qsum['total_files']} 个"),
            ("质量正常文件",   f"{ok_files} 个"),
            ("发现问题文件",   f"{qsum['files_with_issues']} 个"),
            ("存在问题传感器", f"{qsum['sensors_with_issues']} 种"),
            ("交付总数据量",   _human_size(qsum['total_size_mb'])),
        ]
        tbl_q0 = doc.add_table(rows=len(overview_q), cols=2)
        tbl_q0.style = "Table Grid"
        for i, (k, v) in enumerate(overview_q):
            row = tbl_q0.rows[i]
            row.cells[0].text = k
            row.cells[1].text = v
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
        doc.add_paragraph()

        # 逐传感器明细
        if qr_results:
            doc.add_paragraph("各传感器 / 产品质量明细：")
            for sq in qr_results:
                has_any_issue = sq.issues or any(fq.issues for fq in sq.files)
                status_mark = "⚠" if has_any_issue else "✓"
                heading_text = (
                    f"{status_mark}  {sq.season} — {sq.sensor_label}"
                    f"  ({len(sq.files)} 个文件)"
                )
                h_para = doc.add_paragraph()
                h_run = h_para.add_run(heading_text)
                h_run.bold = True
                h_run.font.size = Pt(10)

                # 文件明细表
                if sq.files:
                    col_headers = ["文件名", "大小", "尺寸(px)", "分辨率(m)",
                                   "波段数", "nodata%", "统计元数据", "状态"]
                    tbl_f = doc.add_table(rows=1 + len(sq.files), cols=len(col_headers))
                    tbl_f.style = "Table Grid"
                    _add_table_header(tbl_f, col_headers)
                    for ri, fq in enumerate(sq.files, start=1):
                        row = tbl_f.rows[ri]
                        row.cells[0].text = fq.path.name
                        row.cells[1].text = _human_size(fq.size_bytes / (1024 * 1024))
                        row.cells[2].text = (f"{fq.width}×{fq.height}"
                                             if fq.width else "—")
                        row.cells[3].text = f"{fq.res_x_m:.1f}" if fq.res_x_m > 0 else "—"
                        row.cells[4].text = str(fq.bands) if fq.bands else "—"
                        row.cells[5].text = (f"{fq.nodata_ratio:.0%}"
                                             if fq.nodata_ratio > 0 else "0%")
                        row.cells[6].text = "有" if fq.has_statistics else "无"
                        row.cells[7].text = "正常" if not fq.issues else "异常"
                    doc.add_paragraph()

                # ── 问题汇总：共性 vs 个性 ───────────────────────────
                # 用 issue_key（不含数值）统计各类问题跨文件出现次数
                # key → (代表性issue文本, fix文本, 出现次数)
                key_data: Dict[str, List] = {}   # key -> [issue_text, fix_text, count]
                n_files = max(len(sq.files), 1)

                # 传感器级问题（无 key，直接按全文合并）
                for idx, issue in enumerate(sq.issues):
                    fix = sq.fixes[idx] if idx < len(sq.fixes) else "—"
                    if issue not in key_data:
                        key_data[issue] = [issue, fix, 0]
                    key_data[issue][2] += n_files  # 权重=全部文件数，确保归共性

                # 文件级问题：按 issue_key 归并
                for fq in sq.files:
                    for j, key in enumerate(fq.issue_keys):
                        issue_text = fq.issues[j]
                        fix_text = fq.fixes[j] if j < len(fq.fixes) else "—"
                        if key not in key_data:
                            key_data[key] = [issue_text, fix_text, 0]
                        key_data[key][2] += 1

                if key_data:
                    threshold = max(2, math.ceil(n_files * 0.5))
                    common = [(d[0], d[1]) for d in key_data.values() if d[2] >= threshold]
                    unique_keys = {k for k, d in key_data.items() if d[2] < threshold}

                    # 个性问题：带具体文件名
                    unique_issues: List[Tuple[str, str]] = []
                    for fq in sq.files:
                        for j, key in enumerate(fq.issue_keys):
                            if key in unique_keys:
                                fix = fq.fixes[j] if j < len(fq.fixes) else "—"
                                unique_issues.append((f"[{fq.path.name}] {fq.issues[j]}", fix))

                    all_rows: List[Tuple[str, str]] = []
                    if common:
                        all_rows.append(("【共性问题（全部/多数文件）】", ""))
                        all_rows.extend(common)
                    if unique_issues:
                        all_rows.append(("【个别文件特有问题】", ""))
                        all_rows.extend(unique_issues)

                    tbl_i = doc.add_table(rows=1 + len(all_rows), cols=2)
                    tbl_i.style = "Table Grid"
                    _add_table_header(tbl_i, ["发现问题", "修复方向"])
                    for ri, (issue, fix) in enumerate(all_rows, start=1):
                        tbl_i.rows[ri].cells[0].text = issue
                        tbl_i.rows[ri].cells[1].text = fix
                        if issue.startswith("【"):
                            for run in tbl_i.rows[ri].cells[0].paragraphs[0].runs:
                                run.bold = True
                    doc.add_paragraph()

        else:
            doc.add_paragraph("（delivery 目录中未找到可分析的 GeoTIFF 文件）")

        # 全局质量结论
        doc.add_paragraph()
        if qsum["files_with_issues"] == 0:
            conclusion = "本次交付所有文件质量检查通过，未发现异常。"
        else:
            conclusion = (
                f"本次交付共发现 {qsum['files_with_issues']} 个文件存在质量问题，"
                f"涉及 {qsum['sensors_with_issues']} 种传感器，"
                "请按上表修复方向逐项排查，在下次打包时重新生成报告验证。"
            )
        p_conc = doc.add_paragraph(conclusion)
        p_conc.runs[0].bold = True

    except Exception as _qe:
        doc.add_paragraph(f"（质量分析执行异常：{_qe}）")

    # ── 保存 ─────────────────────────────────────────────────────
    try:
        doc.save(report_path)
        print(f"  [报告] 已生成: {report_name}")
        return report_path
    except Exception as e:
        print(f"  [警告] 报告保存失败: {e}")
        return None
